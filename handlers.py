import customtkinter as ctk
import os 
import re
import string
import traceback
import threading
from docx import Document
import config
from Frames import root,file_path,home_img,create_img,exp_img,content_frame,home_frame,create_frame,example_frame,algo_example,basic1_menu,basic2_menu,norm_menu,main_frame,tab_frame
from file_func import browse_folder,create_file,browse_file,rebuild_cpp_code,rebuild_pascal_code,rebuild_python_code,rebuild_java_code,run_code_program
from Data import Combo_list,value_of_array_in_test,Checkbox_list,Func_limit_list,FuncItem,Func_list,SO_PHAN_TU,KI_TU,MANG_KI_TU
from content_menu import display_content, cons_click
from func_ran import Do_dai_phan_tu,Mang_so,ran_num,Ki_tu,Mang_ki_tu,So_bo_truy_van   
from Back_end_version2 import generate_random_numbers,random_string,random_number,generate_graph,random_edge_count
list1 = ["0","1","2","3","4","5","6","7","8","9"
     ,"a","b","c","d","e","f","g","h","i","j","k","l","m","n","o","p","q","r","s","t","u","v","w","x","y","z"
     ,"A","B","C","D","E","F","G","H","I","J","K","L","M","N","O","P","Q","R","S","T","U","V","W","X","Y","Z"
     ,"!","@","#","$","%","^","&","*","(",")",",",".","?","/","\\","|"
    ]
shared_graph = [0]


current_pop = None
def run_functions(create_test_list):
    global Func_limit_list
    global shared_graph
    results = []  # To store the results of each function run
    t = 0
    for item in create_test_list:
        func = item[0].ham  # The function to run
        func_name = item[0].name  # The name of the function for identifying parameters
        # print(cr)
        # Example function calls based on the function name
        if(t != 0):
            t = t - 1
            continue
        if func_name == "Số phân tử (độ dài)":
            if(item[0].hide_or_show == True):
                if(shared_graph[0] == 0):
                    results.append([item[0].gt,item[0].end_line])
                else: 
                    results.append([str(item[0].gt) +" "+ str(item[0].gt_graph[0]),item[0].end_line])     
        elif func_name == "Số bộ truy vấn":
            results.append([item[0].gt,item[0].end_line])
        elif func_name == "Bắt đầu truy vấn":
            result = ""
            for k in range(item[0].loop):
                gt1 = item[0].vtq_start 
                gt2 = item[0].vtq_end 
                result = result + (tao_test(Combo_list[gt1 + 1 : gt2] , Func_limit_list[gt1 + 1: gt2 ]))
                # print(result)
            t = gt2 - gt1
            item[0].gt = result
            results.append([result,item[0].end_line])
        elif func_name == "Mảng số (nguyên/thực)":
            # Assuming the function expects positional arguments
            result = func(
                min_value = item[0].limit[0],
                max_value = item[0].limit[1],  # Positional arguments instead of keyword
                loop = item[0].loop,
                decimal_places = item[0].decimal_places,
                unique = item[0].unique,
                ascending= item[0].ascending,
                descending= item[0].descending,
                odd_only= item[0].odd_only,
                even_only= item[0].even_only,
                no_zero=item[0].no_zero,
                arithmetic_progression=item[0].arithmetic_progression,
                common_diff_range=item[0].common_diff_range
            )
            item[0].gt = result
            results.append([result,item[0].end_line])
        
        elif func_name == "Xâu kí tự":
            # Assuming this requires multi-word, length, spaces, etc.
            result = func(
                multi_word = item[0].multi_word,
                list_value = item[0].limit,
                length = item[0].loop,  # Pass positional arguments
                min_word_length = item[0].min_word_length,
                max_word_length =item[0].max_word_length,
                min_spaces_in_middle = item[0].min_spaces_in_middle,
                max_spaces_in_middle = item[0].max_spaces_in_middle,
                min_spaces_at_start = item[0].min_spaces_at_start,
                max_spaces_at_start = item[0].max_spaces_at_start,
                forbidden_start_chars = item[0].forbidden_start_chars,
                combine_previous = item[0].combine_previous,
                previous_string = item[0].previous_string
            )
            item[0].gt = result
            results.append([result,item[0].end_line])

        elif func_name == "Số ngẫu nhiên":
            # For random numbers, based on the flags (even, odd, prime, etc.)
            result = func(
                min_value =item[0].limit[0], 
                max_value = item[0].limit[1], # Pass positional arguments
                even = item[0].even,
                odd = item[0].odd,
                prime = item[0].prime,
                perfect_square = item[0].perfect_square,
                decimal_places = item[0].decimal_places
            )
            item[0].gt = result
            results.append([result,item[0].end_line])
        elif func_name == "Kí tự":
            result = func(
                VALUE_ARR = item[0].limit
                )
            item[0].gt = result
            results.append([result,item[0].end_line])
        elif func_name == "Mảng kí tự":
            result = func(
                limit = item[0].limit,
                LOOP = item[0].loop
                )
            item[0].gt = result
            results.append([result,item[0].end_line])
        elif func_name == "Đồ thị":
            result = func(
                num_vertices = item[0].num_vertices,
                num_edges = item[0].num_edges[0],
                directed = item[0].directed,
                allow_self_loops = item[0].allow_self_loops,
                allow_repeated_edges = item[0].allow_repeated_edges,
                connected = item[0].connected,
                weighted = item[0].weighted,
                min_weight = item[0].min_weight,
                max_weight = item[0].max_weight,
                graph_type = item[0].graph_type
                )
            results.append([result,item[0].end_line])
        elif func_name == "Xuống dòng":
            item[0].gt = ""
            results.append(["\n",item[0].end_line])
        # Add more elif cases here for other function types as needed.
    # print(results)
    return results





def func_to_num(n):
    if(n == "Xuống dòng"): return FuncItem(name="Xuống dòng", gt = "\n");
    if(n == "Số phân tử (độ dài)" ): return FuncItem(name="Số phân tử (độ dài)",end_line = None, limit=None, gt=None,gt_graph = None,hide_or_show = True, ham=SO_PHAN_TU, type = "int");
    if(n == "Mảng số (nguyên/thực)" ): return FuncItem(name="Mảng số (nguyên/thực)", limit=None, loop=None, decimal_places=None, unique=None, 
             ascending=None, descending=None, odd_only=None, even_only=None,end_line = None, no_zero = None,
             common_diff_range = None , arithmetic_progression = None,
             ham=generate_random_numbers, gt=[],type = "int");
    if(n == "Số ngẫu nhiên"): return FuncItem(name="Số ngẫu nhiên",end_line = None, limit=None, even=None, odd=None, perfect_square=None, prime=None, decimal_places = None,
             ham=random_number, gt=None,type = "int");
    if(n == "Xâu kí tự"): return FuncItem(name="Xâu kí tự", end_line = None,limit = None,length = None , ham = random_string ,min_word_length = None , max_word_length = None, 
            min_spaces_in_middle = None, max_spaces_in_middle = None, min_spaces_at_start = None ,max_spaces_at_start = None, forbidden_start_chars = None ,
            multi_word = None, combine_previous = None, previous_string = None,type = "string");
    if(n == "Kí tự"): return FuncItem(name="Kí tự", end_line = None,limit=None, ham=KI_TU, gt="",type = "string");
    if(n == "Mảng kí tự"): return FuncItem(name="Mảng kí tự", end_line = None,limit=None, loop=None, ham=MANG_KI_TU, gt=[],type = "string");
    if(n == "Số bộ truy vấn"): return FuncItem(name="Số bộ truy vấn", end_line = None,limit=None, ham = SO_PHAN_TU, type = "int");
    if(n == "Bắt đầu truy vấn"): return FuncItem(name="Bắt đầu truy vấn", vtq_start = None,vtq_end = None, loop = None, gt = None);
    if(n == "Kết thúc truy vấn"): return FuncItem(name="Kết thúc truy vấn", vtq_end = None);
    if(n == "Đồ thị"): return FuncItem(name="Đồ thị",end_line = None,ham = generate_graph,min_num_edges = None, max_num_edges = None,num_vertices = None, num_edges = None, directed = None, allow_self_loops = None, allow_repeated_edges = None,connected = None,weighted = None,min_weight = None, max_weight = None, graph_type = None,gt = None, type = "int")






# Gán dữ liệu đầu vào cho các hàm ( Ràng buộc, giới hạn)
def tao_test(Combo_list,Func_limit_list):
    global shared_graph
    output_string = ""
    create_test_list = []
    nums_of_arr = []
    Func_limit_list_tt = []
    for i,combo in enumerate(Combo_list):
        func_item = func_to_num(combo[0].get())
        create_test_list.append([func_item,combo[1],combo[2]])
        create_test_list[i][0].end_line = combo[3].get()
        # print(combo[3].get())\
        # if(func_item.name == "Đồ thị"):
        #     print(combo[1])
        Func_limit_list_tt.append(Func_limit_list[i])
        
    #xuly giới hạn của
    temp1 = 0
    temp2 = 0
    temp3 = 0
    for i, func_item in enumerate(create_test_list):
        limit = Func_limit_list_tt[i].get().split()  # Initialize default values for limits
        # print(limit)
        if(func_item[0].name == "Bắt đầu truy vấn"):
            func_item[0].vtq_start = i
            temp1 = i
        if(func_item[0].name == "Kết thúc truy vấn"):
            func_item[0].vtq_end = i
            temp2 = i

        if(func_item[0].type == "int"):
            func_item[0].limit = [int(limit[0]),int(limit[1])]
        if(func_item[0].type == "string"):
            func_item[0].limit = limit
    # print(temp1,temp2)
    for item in create_test_list:
        # print(item[0].name)
        if(item[0].name == "Số phân tử (độ dài)" or item[0].name == "Số bộ truy vấn" ):
            if(item[0].limit is not None and item[0].limit[0] <= item[0].limit[1] ):
                item[0].gt = item[0].ham(item[0].limit)
        if(item[0].name == "Số bộ truy vấn"):
            temp3 = item[0].gt
    #Vòng lặp cho các đối tượng yêu cầu đối 
    for item in create_test_list:
        if(item[0].name in ["Mảng số (nguyên/thực)","Xâu kí tự","Mảng kí tự","Đồ thị"]):
            d = 0
            for item2 in create_test_list:  
                if(item2[0].name == "Số phân tử (độ dài)"):
                    d+= 1
                if(item[2] == item2[0].name +" "+ str(d) and item[0].name == "Đồ thị"):
                    item[0].num_vertices = item2[0].gt
                    item[0].min_num_edges = item[0].limit[0]
                    item[0].max_num_edges = item[0].limit[1]
                    item2[0].gt_graph = shared_graph
                    item[0].num_edges = shared_graph
                elif(item[2] == item2[0].name +" "+ str(d)):
                    item[0].loop = item2[0].gt
                
        # if(item[0].name in ["Đồ thị","Ma trận"]):
        #     d = 0
        #     for item2 in create_test_list:
        #         if(item2[0].name == "Số phân tử (độ dài)"):
        #             d+= 1
        #         if(item[2][0] == item2[0].name +" "+ str(d)):
        #             item[0].limit1 = item2[0].limit
        #         if(item[2][1] == item2[0].name +" "+ str(d)):
        #             item[0].limit2 = item2[0].limit
        if(item[0].name == "Bắt đầu truy vấn"):
            item[0].vtq_start = temp1
            item[0].vtq_end = temp2
            item[0].loop = temp3
        if(item[1] is not None):
            if(item[0].name == "Số phân tử (độ dài)"):
                print(item[1][0])
                item[0].hide_or_show = item[1][0]
            if(item[0].name == "Mảng số (nguyên/thực)"):
                item[0].unique = item[1][1]
                item[0].ascending = item[1][2]
                item[0].descending = item[1][3]
                item[0].even_only = item[1][4]
                item[0].odd_only = item[1][5]
                item[0].no_zero = item[1][6]
                item[0].arithmetic_progression = item[1][7]
                item[0].common_diff_range = (item[1][8],item[1][9])
                if(item[1][10] == 1):
                    item[0].decimal_places = item[1][11]
                else:
                    item[0].decimal_places = None
            if(item[0].name == "Số ngẫu nhiên"):
                if(item[1][0] == 1):
                    item[0].decimal_places = item[1][1]
                else:
                    item[0].decimal_places = None
                item[0].even = item[1][3]
                item[0].odd = item[1][2]
                item[0].prime = item[1][4]
                item[0].perfect_square = item[1][5]
            if(item[0].name == "Xâu kí tự"):
                if(item[1][0] == 1):
                    item[0].multi_word = item[1][0]
                    item[0].min_word_length = item[1][1][0]
                    item[0].max_word_length = item[1][1][1]
                    item[0].min_spaces_in_middle = item[1][2][0]
                    item[0].max_spaces_in_middle = item[1][2][1]
                    item[0].min_spaces_at_start = item[1][3][0]
                    item[0].max_spaces_at_start = item[1][3][1]
                if(item[1][4][0] == 1):
                    item[0].forbidden_start_chars = item[1][4][1]
                item[0].combine_previous = item[1][5]
                if(item[0].combine_previous == 1):
                    item[0].previous_string = ""
                else:
                    item[0].previous_string = "\n"
            if(item[0].name == "Đồ thị"):
                item[0].directed = item[1][0]
                item[0].allow_self_loops = item[1][1]
                item[0].allow_repeated_edges = item[1][2]
                item[0].connected = item[1][3]
                item[0].weighted = item[1][4]
                item[0].min_weight = item[1][5][0]
                item[0].max_weight = item[1][5][1]
                item[0].graph_type = item[1][6]
                shared_graph[0] = random_edge_count(num_vertices = item[0].num_vertices , min_edges = item[0].min_num_edges , max_edges = item[0].max_num_edges, directed = item[0].directed, graph_type = item[0].graph_type)
        elif(item[0].name == "Đồ thị"):
            shared_graph[0] = random_edge_count(num_vertices = item[0].num_vertices , min_edges = 0 , max_edges = item[0].num_vertices, directed = False)
                # print(shared_graph[0])
    # for item in create_test_list:
    res = run_functions(create_test_list)
    for gt in res:
        if isinstance(gt[0], list):
            if gt[1] == True:
                for value in gt[0]:
                    if isinstance(value, list):
                        for value2 in value:
                            output_string += f"{value2} "
                        output_string+="\n"
                    else:
                        output_string += f"{value}\n"  # Newline for each value
            else:
                for value in gt[0]:
                    output_string += f"{value} "  # Space-separated values
        else:
            if gt[1] == True:
                output_string += f"{gt[0]}\n" 
            else:
                output_string +=f"{gt[0]}" 
    return output_string





def select_checkbox_all_type(limit1,limit2,value):
    if value == 1:
        for i in range(limit1,limit2+1):
            Checkbox_list[i].select()
    if value == 0:
        for i in range(limit1,limit2+1):
            Checkbox_list[i].deselect()

def click_get_limit(pop,min_value,max_value,entry_space):
    entry_space.delete(0,ctk.END)
    flag_max = 0  
    flag_min = 0
    flag_char = 0
    if(min_value.get() != ""):
        flag_min = 1
    if(max_value.get() != ""): 
        flag_max = 1
    for i in range (len(Checkbox_list)):
        if(Checkbox_list[i].get() == True):
            flag_char = 1
            break
        # if()
    # print("Gay")
    if(flag_min + flag_max <= 2 and flag_min + flag_max != 0):
        temp_min = min_value.get()
        temp_max = max_value.get()
        if(flag_min + flag_max == 1):
            if(flag_min == 0):
                temp_min = temp_max
            else:
                temp_max = temp_min
        entry_space.insert(0,temp_min + " " + temp_max)
        pop.destroy()
    if(flag_char == 1):
        check_list_limit = []
        for i in range (len(Checkbox_list)):
            if(Checkbox_list[i].get() == True):
                check_list_limit.append(list1[i])
        entry_space.insert(0,check_list_limit)
        pop.destroy() 


def limit_click(entry_space):
    global current_pop
    if current_pop is not None and current_pop.winfo_exists():
        current_pop.destroy()

    # Tạo hộp thoại chính
    pop = ctk.CTkToplevel()
    pop.title("Hộp thoại giới hạn")
    pop.geometry("525x625")
    pop.resizable(False, False)
    pop.transient(root)
    pop.lift()
    pop.focus_force()
    current_pop = pop
    Checkbox_list.clear()

    # Header frame
    header_frame = ctk.CTkFrame(pop, corner_radius=10)
    header_frame.grid(row=0, column=0, columnspan=2, pady=10, padx=10, sticky="nsew")

    header_label = ctk.CTkLabel(header_frame, text="Thiết lập giới hạn", font=("Arial", 20, "bold"))
    header_label.pack(pady=10)

    # ComboBox frame
    combo_frame = ctk.CTkFrame(pop, corner_radius=10)
    combo_frame.grid(row=1, column=0, columnspan=2, pady=10, padx=10, sticky="nsew")

    # Thêm label cho frame
    combo_label = ctk.CTkLabel(combo_frame, text="Giới hạn cho dữ liệu số", font=("Arial", 15, "bold"))
    combo_label.grid(row=0, column=1, columnspan=3, pady=10)
    # Negative ComboBox
    min_label = ctk.CTkLabel(combo_frame, text="Giới hạn nhỏ nhất:", font=("Arial", 13))
    min_label.grid(row=1, column=0, padx=10, pady=5, sticky="e")
    
    # numbs = ["Số âm"] + [str(-10**i // 10) for i in range(10, 1, -1)]
    min_value = ctk.CTkEntry(combo_frame)
    min_value.grid(row=1, column=1, padx=10, pady=5, sticky="w")

    # Positive ComboBox
    max_label = ctk.CTkLabel(combo_frame, text="Giới hạn lớn nhất:", font=("Arial", 13))
    max_label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
    
    # res = ["Số dương"] + [str(10**i // 10) for i in range(10, 1, -1)]
    max_value = ctk.CTkEntry(combo_frame)
    max_value.grid(row=2, column=1, padx=10, pady=5, sticky="w")

    # Checkbox frame
    checkbox_frame = ctk.CTkFrame(pop, corner_radius=10)
    checkbox_frame.grid(row=2, column=0, columnspan=2, pady=10, padx=10, sticky="nsew")

    # Shortcut checkboxes
    shortcut_frame = ctk.CTkFrame(checkbox_frame)
    shortcut_frame.grid(row=0, column=0, columnspan=10, pady=10, padx=10)

    checkbox_name = ctk.CTkLabel(shortcut_frame, text="Giới hạn cho dữ liệu ký tự", font=("Arial", 15, "bold"))
    checkbox_name.grid(row=0, column=0, columnspan=4, pady=5)

    temp = [
        {"name": "Chữ số", "limit1": 0, "limit2": 9},
        {"name": "Ký tự thường", "limit1": 10, "limit2": 35},
        {"name": "Ký tự hoa", "limit1": 36, "limit2": 61},
        {"name": "Ký tự đặc biệt", "limit1": 62, "limit2": 77},
    ]

    for i, item in enumerate(temp):
        shortcut_box = ctk.CTkCheckBox(shortcut_frame, text=item["name"], checkbox_width=10, checkbox_height=10)
        shortcut_box.configure(command=lambda i=i, box=shortcut_box: select_checkbox_all_type(temp[i]["limit1"], temp[i]["limit2"], box.get()))
        shortcut_box.grid(row=1, column=i, padx=10, pady=5)

    # Dynamic checkboxes
    n = 2
    row_counter = 2
    current_section = None

    for char in list1:
        if char in string.punctuation:
            section = 1
            val = "#eb4034"
        elif char in string.digits:
            section = 2
            val = "#83eb34"
        elif char in string.ascii_lowercase:
            section = 3
            val = "#3477eb"
        elif char in string.ascii_uppercase:
            section = 4
            val = "#ebb134"
        else:
            section = None

        if section != current_section:
            current_section = section
            n = 0
            row_counter += 1

        checkbox = ctk.CTkCheckBox(
            checkbox_frame,
            text_color=val,
            text=char,
            checkbox_width=10,
            checkbox_height=10,
            width=50,
            height=10
        )
        Checkbox_list.append(checkbox)
        checkbox.grid(row=row_counter, column=(n % 10), pady=5)

        n += 1
        if n % 10 == 0:
            row_counter += 1

    # OK button
    get_data_button_limit = ctk.CTkButton(pop, text="Xác nhận", width = 250, height = 50,font=("Arial", 20, "bold"),border_color = "white",border_width= 2,
                                          command=lambda entry_space=entry_space: click_get_limit(pop = pop, min_value = min_value, max_value = max_value, entry_space = entry_space))
    get_data_button_limit.grid(row=3, column=0, columnspan=2, pady=10)




 
    # def confirm_selection():
    #     selected = []
    #     for checkbox in Checkbox_list:
    #         if checkbox.get() == 1:  
    #             selected.append(checkbox.cget("text"))
    #     label.config(text="Selected: " + ", ".join(selected))
    #     pop.destroy()   

    # confirm_button = ctk.CTkButton(pop, text="Confirm", command=confirm_selection)
    # confirm_button.grid(row=row_counter + 1, column=0, columnspan=10, pady=10)






def create_buttons(frame, button_texts, target):
    buttons=[]
    
    temp=[]
    def on_button_click(num): 
        temp = config.problems.get(button_texts[num-1])   
        def create2(frame1, button_texts):
            buttons1=[] 
            for widget in frame1.winfo_children():
                widget.destroy()

            for index, text in enumerate(button_texts, start=1):
                button = ctk.CTkButton(frame1, text=text, width=200, command = lambda num = index: display_content(frame1, button_texts[num-1]))
                button.pack(pady=10, padx=10)
                buttons1.append(button)
            return buttons1

        if(temp):
            # print(temp) 
            create2(target, temp) 
        if(button_texts==config.math):
            display_content(target, config.math[num-1])
    
    
    if temp: 
        # print(temp)
        button_texts=temp
    for index, text in enumerate(button_texts, start=1):

        button = ctk.CTkButton(frame, text=text, width=200,
                               command=lambda num=index: (on_button_click(num), target.tkraise()))   
        button.pack(pady=10, padx=10)
        buttons.append(button)
    
    close_button = ctk.CTkButton(target,
                                width=70, height=20,
                                text="X",
                                command=frame.tkraise(),)
    close_button.pack(padx=10, pady=10, anchor="e")

    return buttons

def combo_click(frame, value, row, widget_rows):  
    widgets_to_hide = widget_rows.get(row)
    if widgets_to_hide:
        if value in ["Xuống dòng", "Kết thúc truy vấn", "Bắt đầu truy vấn"]:
            # Ensure you're not going out of bounds when hiding widgets
            for widget in widgets_to_hide[1:3]: 
                widget.grid_remove()
        # elif value in ["Xâu kí tự", "Kí tự"]:
        #     # Hide a specific range of widgets (only widget 3 in this case)
        #     for widget in widgets_to_hide[3:4]:  
        #         widget.grid_remove()
        else:
            for widget in widgets_to_hide[1:]:
                widget.grid()

    if not value in ["Mảng kí tự", "Mảng số (nguyên/thực)","Xâu kí tự","Đồ thị"]:
        for widget in widgets_to_hide[4:]:
            widget.grid_remove()
    if value == "Mảng kí tự":
        for widget in widgets_to_hide[3:4]:
            widget.grid_remove()

    for i,combo in enumerate(Combo_list):
        widget_temp = widget_rows.get(i+1)
        if widget_temp:
            # widget_temp[4].configure(values = value_of_array_in_test)
            if(combo[0].get() == "Mảng kí tự" or combo[0].get() == "Mảng số (nguyên/thực)" or combo[0].get() == "Xâu kí tự" or combo[0].get() == "Đồ thị"):     
                d = 0
                temp = ["0"]
                for combo2 in Combo_list[0:i]:
                    if(combo2[0].get() ==  "Số phân tử (độ dài)" ):
                        d = d + 1
                        temp.append("Số phân tử (độ dài) " + str(d))
                    global value_of_array_in_test
                    value_of_array_in_test = temp
                    widget_temp[4].configure(values = value_of_array_in_test)
                    widget_temp[4].grid_remove()
                    widget_temp[4].grid()

def get_size_of_data(selected_va,row):
    Combo_list[row-1][2] =(selected_va)

def add_inputs(create_frame, row_counter, widget_rows, val, Combo_list):
    combo_new = ctk.CTkOptionMenu(create_frame, values=val, 
                                  command=lambda selected_value: combo_click(create_frame, selected_value, row_counter, widget_rows))
    combo_new.grid(row=row_counter, column=0, padx=10, sticky="w") 

    cons_new = ctk.CTkEntry(create_frame, width=150, height=10)
    cons_new.grid(row=row_counter, column=1, sticky="w")
    cons_new.bind("<Button-1>", lambda event: cons_click(row_counter, combo_new.get()))

    limit1_new = ctk.CTkEntry(create_frame, width=120, height=10)
    limit1_new.grid(row=row_counter, column=2, padx=10, pady=10, sticky="w")
    limit1_new.bind("<Button-1>", lambda event: limit_click(limit1_new))
    
    switch_var = ctk.IntVar(value=1)  # Default value is checked
    switch = ctk.CTkCheckBox(
        create_frame,
        text="xuống dòng",
        variable=switch_var
    )
    switch.grid(row=row_counter, column=3, pady=10, sticky="w")
    # limit2_new = ctk.CTkEntry(create_frame, width=120, height=10)
    # limit2_new.grid(row=row_counter, column=3, pady=10, sticky="w")
    # limit2_new.bind("<Button-1>", lambda event: limit_click(limit2_new))
    global value_of_array_in_test
    # print(value_of_array_in_test)
    option = ctk.CTkOptionMenu(create_frame, values=value_of_array_in_test, command = lambda selected_value : get_size_of_data(selected_value,row_counter))
    # Add the option menu, even if hidden initially
    if row_counter>=1:
       option.grid_configure(row=row_counter, column=6, padx=10, sticky="w")
       option.grid_remove()
    
    # Add all widgets to Combo_list and widget_rows
    Combo_list.append([combo_new,None,None,switch])
    Func_limit_list.append((limit1_new))
    widget_rows[row_counter] = (combo_new, cons_new, limit1_new, switch, option)

def choose_browser_folder(entry):
    entry.delete(0,'end')
    entry.insert(0,browse_folder())
def choose_browser_file(entry):
    entry.delete(0,'end')
    entry.insert(0,browse_file(1))
def create_file_test(
    language=None,
    template=None,
    dir_file_code=None,
    number_of_test=None,
    name_file=None,
    dir_folder=None,
    on_progress=None,  # Callback cập nhật ProgressBar
    on_log=None        # Callback ghi nhật ký vào textbox
):
    def log(message):
        if on_log:
            on_log(message)

    if number_of_test == "" or dir_folder == "":
        log("Thư mục hoặc số lượng test không hợp lệ.")
        return

    flag_file_out = (dir_file_code is not None and dir_file_code != "" and dir_file_code != "none")
    number_of_test = int(number_of_test)
    dir_file_code_exe = re.sub(r'\..*', '', dir_file_code) + ".exe"

    # Tính tổng bước để cập nhật progress bar
    progress_total_steps = 3
    if flag_file_out:
        progress_total_steps = 4
    current_step = 0

    def update_progress():
        nonlocal current_step
        current_step += 1
        if on_progress:
            on_progress(current_step / progress_total_steps)

    if flag_file_out:  # Biên dịch file code nếu cần    
        log(f"Đang biên dịch file: {dir_file_code}")
        if language == "C++":
            rebuild_cpp_code(dir_file_code, dir_file_code_exe)
        elif language == "Pascal":
            rebuild_pascal_code(dir_file_code, dir_file_code_exe)
        elif language == "Python":
            rebuild_python_code(dir_file_code, dir_file_code_exe)
        elif language == "Java":
            rebuild_java_code(dir_file_code, dir_file_code_exe)
        log("Biên dịch thành công.")
        update_progress()

    os.makedirs(dir_folder + name_file, exist_ok=True)
    name_dir_for_test = os.path.join(dir_folder, name_file)
    # TEMPLATE AND GENERATE TEST
    update_progress()
    if flag_file_out: 
        log("Đang khởi tạo input và output")
    else:
        log("Đang khởi tạo input")
    if template == "Xuất Folder":
        for i in range(number_of_test):
            inp = tao_test(Combo_list,Func_limit_list)
            # log(f"Tạo file input: {name_file}{i+1}.INP")
            create_file(name_file + str(i + 1) + ".INP", name_dir_for_test + "\\", inp)
            if flag_file_out and inp:
                res_output = run_code_program(dir_file_code_exe, inp)
                create_file(name_file + str(i + 1) + ".OUT", name_dir_for_test + "\\", res_output)
        log("Tạo dạng Folder thành công")
        update_progress()

    elif template == "Xuất Themis":
        for i in range(number_of_test):
            test_case_folder = os.path.join(name_dir_for_test, f"TEST{i+1}")
            os.makedirs(test_case_folder, exist_ok=True)
            inp = tao_test(Combo_list,Func_limit_list)
            # log(f"Tạo file input: TEST{i+1}/{name_file}.INP")
            create_file(name_file + ".INP", test_case_folder + "\\", inp)
            if flag_file_out and inp:
                # log("Chạy chương trình để tạo output.")
                res_output = run_code_program(dir_file_code_exe, inp)
                create_file(name_file + ".OUT", test_case_folder + "\\", res_output)
        log("Tạo dạng Themis thành công")
        update_progress()

    elif template == "Xuất Word":
        tasks = [name_file]
        inputs =[]
        outputs=[]
        inputs_for_test =[]
        outputs_for_test =[]
        for i in range(number_of_test):   
            inp = tao_test(Combo_list,Func_limit_list)
            inputs_for_test.append(inp)
            if(flag_file_out and (inp != None or inp != "")):
                res_output = run_code_program(dir_file_code_exe,inp)
                outputs_for_test.append(res_output)    
    # Create a new Document
        inputs.append(inputs_for_test)
        outputs.append(outputs_for_test)
        doc = Document()
        for i, task in enumerate(tasks):
            # Add task heading
            doc.add_heading(task, level=2)
            
            # Find the max length of inputs/outputs for rows in the table
            max_len = max(len(inputs[i]), len(outputs[i]))
            
            # Create a table with 2 columns and max_len + 1 rows (first row for headers)
            table = doc.add_table(rows=max_len + 1, cols=2)
            
            # Set column headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Input'
            hdr_cells[1].text = 'Output'
            
            # Fill in input and output data
            for j in range(max_len):
                row_cells = table.rows[j + 1].cells
                row_cells[0].text = inputs[i][j] if j < len(inputs[i]) else ''
                row_cells[1].text = outputs[i][j] if j < len(outputs[i]) else ''
            
            # Add some space between tasks
            doc.add_paragraph("\n")
        doc.save(dir_folder +"\\"+name_file + ".docx")
        update_progress()
        log("Tạo dạng Word thành công.")
    log("\n")
    log("Hoàn tất tạo test cases.")
    log("Thông tin của test bạn vừa tạo:")
    log(f"Ngôn ngữ: {language}")
    log(f"Form: {template}")
    log(f"Số lượng: {number_of_test} ")
    log(f"Tên file: {name_file}" )
    log(f"Tên địa chỉ: {dir_folder}")
    update_progress()


def create_file_test_threaded(
    language=None,
    template=None,
    dir_file_code=None,
    number_of_test=None,
    name_file=None,
    dir_folder=None,
    progress_bar=None,  # ProgressBar từ giao diện
    on_complete=None,   # Callback khi hoàn tất
    textbox=None
):  
    textbox.configure(state = "normal")
    textbox.delete("0.0", "end")
    textbox.insert("1.0", "Log thông báo sẽ xuất hiện tại đây...\n")  # Nội dung mẫu
    def task():
        try:
            def log(message):
                if textbox:
                   textbox.insert("end", message + "\n")
                   textbox.see("end")  # Cuộn xuống dòng cuối

            def update_progress(value):
                if progress_bar:
                    progress_bar.set(value)  # Cập nhật thanh progress

            log("Bắt đầu tạo test cases...")
            create_file_test(
                language=language,
                template=template,
                dir_file_code=dir_file_code,
                number_of_test=number_of_test,
                name_file=name_file,
                dir_folder=dir_folder,
                on_progress=update_progress,
                on_log=log,
            )
            if on_complete:
                on_complete()
            else:
                log("Tạo test hoàn tất!")
                textbox.configure(state = "disabled")
        except Exception as e:
            if textbox:
                tb_info = traceback.format_exc()
                textbox.insert("end", f"Lỗi: {e}\n")
                textbox.insert("end", f"Chi tiết lỗi:\n{tb_info}\n")
                textbox.see("end")
                textbox.configure(state = "disabled")
    thread = threading.Thread(target=task)
    thread.daemon = True
    thread.start()
    

def create_file_test_ex(file_dir,file_name, loop):
    os.makedirs(file_dir + file_name, exist_ok=True) 
    name_dir_for_test  = file_dir + file_name
    for i in range(loop):
        os.makedirs(name_dir_for_test + "\\TEST" + str(i+1))
        name_dir_for_test_case = name_dir_for_test + "\\TEST" + str(i+1)     
        inp = tao_test(None)
        create_file(file_name + ".INP",name_dir_for_test_case + "\\",inp)
        create_file(file_name + ".OUT",name_dir_for_test_case + "\\","xucaliat")

    

def remove_inputs(widget_rows, row_counter):
    if row_counter in widget_rows:
        for widget in widget_rows[row_counter]:
            widget.destroy()   
        del widget_rows[row_counter]